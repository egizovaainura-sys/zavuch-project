import streamlit as st
import pandas as pd
import sqlite3
from datetime import datetime
import plotly.express as px
import io
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- 1. НАСТРОЙКИ СТРАНИЦЫ ---
st.set_page_config(page_title="Smart Завуч: Фокус-группа", layout="wide")

# --- 2. БАЗА ДАННЫХ (Для отчетов) ---
def init_db():
    conn = sqlite3.connect('school_focus_lite.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS reports (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id TEXT, 
        date TEXT, quarter INTEGER, teacher TEXT, student TEXT, subject TEXT, grade TEXT, topic TEXT, goal TEXT,
        purpose TEXT, start_t TEXT, start_s TEXT, middle_t TEXT, middle_s TEXT, end_t TEXT, end_s TEXT,
        ict_usage TEXT, methods TEXT, reflection TEXT,
        reserve_json TEXT, scores_json TEXT, comments_json TEXT,
        s1 TEXT, s2 TEXT, s3 TEXT, g1 TEXT, g2 TEXT, g3 TEXT, advice TEXT, percent REAL, lang TEXT
    )''')
    conn.commit()
    conn.close()

# --- 3. ПРОСТАЯ ПРОВЕРКА ДОСТУПА (БЕЗ СЛОЖНЫХ БИБЛИОТЕК) ---
def check_access_simple(phone_number):
    try:
        # Ваша ссылка на таблицу, преобразованная в формат CSV для скачивания
        # Используем pandas, который уже установлен
        sheet_url = "https://docs.google.com/spreadsheets/d/1Z1BUjdyNm6sv9CvZ-gDmljp2kjhOHWVH7lK-gh53RtQ/export?format=csv"
        
        # Читаем таблицу
        df = pd.read_csv(sheet_url)
        
        # Берем первый столбец и превращаем все номера в строки
        allowed_list = df.iloc[:, 0].astype(str).str.strip().tolist()
        
        # Очищаем ввод пользователя
        clean_phone = str(phone_number).strip()
        
        if clean_phone in allowed_list:
            return True
        return False
    except Exception as e:
        st.error(f"Ошибка проверки доступа. Проверьте интернет или ссылку: {e}")
        return False

# --- 4. СЛОВАРЬ ИНТЕРФЕЙСА ---
LANGS = {
    'RU': {
        'title': "Smart Завуч 🇰🇿", 'header': "ЛИСТ НАБЛЮДЕНИЯ УРОКА (ФОКУС-ГРУППА)",
        'nav_new': "📊 Ввод данных", 'nav_rating': "🏆 Рейтинг", 'nav_map': "📈 Динамика",
        'teacher': "ФИО Учителя", 'student': "ФИО Ученика (Резерв)", 'subject': "Предмет", 'grade': "Класс",
        'date': "Дата", 'quarter': "Четверть", 'topic': "Тема урока", 'goal': "Цели урока",
        'purpose': "Цель посещения", 'res_header': "2. Назардағы оқушылар / Фокус на учащихся 'резерва'",
        'res_fio': "ФИО ученика", 'res_inter': "Взаимодействие учителя",
        'res_react': "Реакция и активность", 'res_idx': "Индекс (УД/ТБ)",
        'crit_header': "3. Общий анализ урока", 'prof_header': "🎯 Профессионализм",
        'ict_label': "Использование ИКТ", 'methods_label': "Методы обучения",
        'reflection': "Рефлексия", 'stages_header': "⏳ Ход урока",
        'conclusion_header': "4. Выводы и рекомендации", 'strengths_label': "Сильные стороны:",
        'growth_label': "Зоны роста:", 'final_advice': "5. Рекомендации учителю",
        'save_btn': "💾 Сохранить", 'excel_btn': "📥 Скачать (Excel)",
        'word_btn': "📄 Скачать (Word)", 'fact_label': "Комментарии",
        'score_label': "Балл", 'action_t': "Действие учителя", 'action_s': "Действие ученика",
        'criteria_list': [
            "Четкость целей", "Содержание материала", "Разнообразие методов", "Дифференциация заданий",
            "Логика этапов", "Критериальное оценивание", "Атмосфера", "Тайм-менеджмент"
        ]
    },
    'KZ': {
        'title': "Smart Завуч 🇰🇿", 'header': "САБАҚТЫ БАҚЫЛАУ ПАРАҒЫ (РЕЗЕРВ)",
        'nav_new': "📊 Деректер енгізу", 'nav_rating': "🏆 Рейтинг", 'nav_map': "📈 Динамика",
        'teacher': "Мұғалімнің АЖТ", 'student': "Оқушының АЖТ", 'subject': "Пән", 'grade': "Сынып",
        'date': "Күні", 'quarter': "Тоқсан", 'topic': "Тақырып", 'goal': "Сабақ мақсаты",
        'purpose': "Бақылау мақсаты", 'res_header': "2. Назардағы оқушылар",
        'res_fio': "Оқушының АЖТ", 'res_inter': "Мұғалімнің әрекеті",
        'res_react': "Оқушының реакциясы", 'res_idx': "Индекстер",
        'crit_header': "3. Жалпы талдау", 'prof_header': "🎯 Кәсіби шеберлік",
        'ict_label': "АКТ қолданылуы", 'methods_label': "Әдіс-тәсілдер",
        'reflection': "Рефлексия", 'stages_header': "⏳ Сабақ кезеңдері",
        'conclusion_header': "4. Қорытынды", 'strengths_label': "Күшті жақтары:",
        'growth_label': "Даму аймақтары:", 'final_advice': "5. Ұсыныстар",
        'save_btn': "💾 Сақтау", 'excel_btn': "📥 Жүктеу (Excel)",
        'word_btn': "📄 Жүктеу (Word)", 'fact_label': "Түсініктеме",
        'score_label': "Баға", 'action_t': "Мұғалім әрекеті", 'action_s': "Оқушы әрекеті",
        'criteria_list': [
            "Мақсаттардың айқындылығы", "Материал мазмұны", "Әдіс-тәсілдер", "Тапсырмаларды саралау",
            "Кезеңдер қисындылығы", "Бағалау", "Психологиялық ахуал", "Уақытты пайдалану"
        ]
    }
}

# --- 5. ФУНКЦИЯ WORD ---
def create_official_docx(data, lang):
    L = LANGS[lang]
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(L['header'])
    run.bold = True
    run.font.size = Pt(14)

    doc.add_heading('1. Информация', level=1)
    doc.add_paragraph(f"{L['teacher']}: {data['teacher']}")
    doc.add_paragraph(f"{L['topic']}: {data['topic']}")
    doc.add_paragraph(f"{L['grade']}: {data['grade']}")
    doc.add_paragraph(f"{L['date']}: {data['date']}")

    doc.add_heading(L['conclusion_header'], level=1)
    doc.add_paragraph(f"РЕКОМЕНДАЦИЯ: {data['advice']}")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 6. ЛОГИКА ---
init_db()

if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False

# ВХОД
if not st.session_state['logged_in']:
    st.title("🔐 Вход для Завуча")
    st.info("Введите номер телефона, который есть в базе (Google Таблица).")
    
    phone_input = st.text_input("Ваш номер:", placeholder="7701xxxxxxx")
    
    if st.button("Войти"):
        if check_access_simple(phone_input):
            st.session_state['logged_in'] = True
            st.session_state['user_id'] = phone_input
            st.success("Успешно!")
            st.rerun()
        else:
            st.error("Номер не найден. Проверьте таблицу.")
    st.stop()

# ПОСЛЕ ВХОДА
st.sidebar.success(f"Вы вошли: {st.session_state['user_id']}")
if st.sidebar.button("Выход"):
    st.session_state['logged_in'] = False
    st.rerun()

lang_choice = st.sidebar.selectbox("Язык", ['RU', 'KZ'])
L = LANGS[lang_choice]

st.title(L['title'])
st.write("Добро пожаловать в рабочую панель!")

# Здесь упрощенная форма для теста (полную можно вернуть позже)
with st.form("main_form"):
    teacher = st.text_input(L['teacher'])
    grade = st.text_input(L['grade'])
    advice = st.text_area(L['final_advice'])
    
    if st.form_submit_button(L['save_btn']):
        st.success("Данные сохранены!")
